#!/usr/bin/env python3
"""
Generate Pending_Tasks_Sprint_Roadmap.docx
Comprehensive task breakdown for Weeks 2-4 sprint execution
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime, timedelta

def add_styled_heading(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.size = Pt(24)
        heading.runs[0].font.color.rgb = RGBColor(184, 150, 12)  # Gold color
    elif level == 2:
        heading.runs[0].font.size = Pt(16)
        heading.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    return heading

def add_table_row(table, cells_data, is_header=False):
    """Add a row to a table"""
    row = table.add_row()
    for i, data in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(data)
        if is_header:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(184, 150, 12)

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('DPCA - Dream Paris Communication Assistant', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_heading('Pending Tasks Sprint Roadmap (Weeks 2-4)', level=2)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(14)
    
    # Date
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.runs[0].font.italic = True
    
    doc.add_paragraph()  # Spacing
    
    # Executive Summary
    add_styled_heading(doc, '📋 Executive Summary', level=1)
    summary_items = [
        'Total Pending Work: ~103 hours across 3 weeks',
        'Team: 3 full-time members (Abdur, Usama, Sophie) + optional DevOps',
        'Critical Path: VPS Infrastructure → n8n Deployment → Workflow Implementation',
        'Critical Dependency: Brand Voice Content (blocks prompt finalization)',
        'Target Completion: End of Week 4 (April 25, 2026)',
        'Status: All Week 1 deliverables 100% complete, Week 2-4 ready to commence'
    ]
    for item in summary_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    # Timeline Overview
    add_styled_heading(doc, '📅 Timeline Overview', level=1)
    week2_date = (datetime.now() + timedelta(days=7)).strftime("%b %d")
    week3_date = (datetime.now() + timedelta(days=14)).strftime("%b %d")
    week4_date = (datetime.now() + timedelta(days=21)).strftime("%b %d")
    
    timeline_items = [
        f'Week 2 ({week2_date}): Infrastructure + VPS setup begins (Abdur)',
        f'Week 3 ({week3_date}): n8n workflows implementation (Abdur) + UI enhancements (Usama) in parallel',
        f'Week 4 ({week4_date}): KB finalization + workflow testing + UI polish + production readiness'
    ]
    for item in timeline_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    # Team Assignments
    add_styled_heading(doc, '👥 Team Assignments & Hours', level=1)
    
    assignments_table = doc.add_table(rows=1, cols=4)
    assignments_table.style = 'Light Grid Accent 1'
    add_table_row(assignments_table, ['Team Member', 'Role', 'Hours', 'Status'], is_header=True)
    add_table_row(assignments_table, ['Abdur Rehman', 'Infrastructure Lead', '45 hrs', 'Ready to start'])
    add_table_row(assignments_table, ['Usama Khan', 'Frontend Lead', '18 hrs', 'Ready to start (parallel)'])
    add_table_row(assignments_table, ['Sophie Laurent', 'Content Lead', '30 hrs', 'CRITICAL - start immediately'])
    add_table_row(assignments_table, ['VPS/DevOps (Optional)', 'Infrastructure Support', '10 hrs', 'Optional - post-MVP'])
    add_table_row(assignments_table, ['TOTAL', '', '103 hours', 'Weeks 2-4'])
    
    doc.add_paragraph()
    
    # ==================== ABDUR REHMAN TASKS ====================
    add_styled_heading(doc, "⚙️ ABDUR REHMAN - Infrastructure & Workflows (45 hours)", level=1)
    
    # Phase 1: VPS Setup
    add_styled_heading(doc, 'Phase 1: VPS Infrastructure Setup (5 hours)', level=2)
    
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Light Grid Accent 1'
    add_table_row(table1, ['Task', 'Description', 'Hours', 'Week'], is_header=True)
    
    abdur_phase1_tasks = [
        ['1.1 VPS Provisioning', 'Rent 2GB RAM server on DigitalOcean/Hetzner (Ubuntu 22.04)', '1 hr', 'Week 2'],
        ['1.2 Docker Setup', 'Install Docker + Docker Compose, configure UFW firewall', '1.5 hrs', 'Week 2'],
        ['1.3 n8n Deployment', 'Deploy n8n via docker-compose.yml (basic auth, persistent volume)', '1.5 hrs', 'Week 2'],
        ['1.4 n8n Config', 'Configure n8n timezone (Europe/Paris), webhook URL, basic auth credentials', '1 hr', 'Week 2'],
    ]
    
    for task in abdur_phase1_tasks:
        add_table_row(table1, task)
    
    doc.add_paragraph('✓ Deliverable: n8n running on http://<VPS_IP>:5678 (accessible + verified)', 
                      style='List Bullet')
    doc.add_paragraph()
    
    # Phase 2: Third-party Credentials
    add_styled_heading(doc, 'Phase 2: OAuth & API Credentials (6 hours)', level=2)
    
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Light Grid Accent 1'
    add_table_row(table2, ['Task', 'Description', 'Hours', 'Week'], is_header=True)
    
    abdur_phase2_tasks = [
        ['2.1 Gmail OAuth', 'Complete Google Cloud OAuth app setup + get credentials (Client ID, Secret)', '2 hrs', 'Week 2'],
        ['2.2 n8n Credentials', 'Create n8n credential nodes: Supabase, OpenAI API, Pinecone, Gmail', '2 hrs', 'Week 2'],
        ['2.3 Meta Setup', 'Apply for Meta App Review (WhatsApp + Instagram APIs) - starts clock for approval', '1 hr', 'Week 2'],
        ['2.4 Backup Creds', 'Store all credentials securely in 1Password or similar vault', '1 hr', 'Week 2'],
    ]
    
    for task in abdur_phase2_tasks:
        add_table_row(table2, task)
    
    doc.add_paragraph('✓ Deliverable: All 5 credential nodes active in n8n + Meta App Review submitted', 
                      style='List Bullet')
    doc.add_paragraph('⚠️ Blocker Note: Meta App Review takes 1-4 weeks. WF1-5 testable without Meta APIs.', 
                      style='List Paragraph')
    doc.add_paragraph()
    
    # Phase 3: n8n Workflows Implementation
    add_styled_heading(doc, 'Phase 3: n8n Workflow Implementation (30 hours)', level=2)
    
    table3 = doc.add_table(rows=1, cols=4)
    table3.style = 'Light Grid Accent 1'
    add_table_row(table3, ['Workflow', 'Description', 'Hours', 'Week'], is_header=True)
    
    workflows = [
        ['WF1: Email Ingestion', 'Gmail polling (2 min interval) → deduplication → Supabase insert', '3 hrs', 'Week 3'],
        ['WF2: Classification', 'GPT-4o message classification (category/priority/tier) + safety override', '3.5 hrs', 'Week 3'],
        ['WF3: KB Embedding', 'text-embedding-3-small (1536-dim) + Pinecone upsert', '2.5 hrs', 'Week 3'],
        ['WF4: Context Retrieval', 'Semantic search from Pinecone (top-5 results) + metadata filtering', '2.5 hrs', 'Week 3'],
        ['WF5: Draft Generation', 'GPT-4o with P1 brand voice system prompt + tone confidence calculation', '4 hrs', 'Week 3-4'],
        ['WF6: Auto-Send', 'Multi-channel routing (Gmail/WhatsApp/Instagram) + retry logic', '4 hrs', 'Week 4'],
        ['WF7: Lead Extraction', 'Structured data extraction via GPT-4o + CRM insert/update to Supabase', '3.5 hrs', 'Week 4'],
        ['WF8: Dashboard Actions', 'Webhook receiver for approve/reject/regenerate/reassign + audit logging', '2 hrs', 'Week 4'],
    ]
    
    for wf in workflows:
        add_table_row(table3, wf)
    
    doc.add_paragraph('✓ Deliverable: All 8 workflows fully implemented + end-to-end pipeline tested', 
                      style='List Bullet')
    doc.add_paragraph('✓ Verification: Run Step 14 checklist from todo.md (all 14 validation points)', 
                      style='List Bullet')
    doc.add_paragraph()
    
    # Phase 4: Testing & Verification
    add_styled_heading(doc, 'Phase 4: Integration Testing (4 hours)', level=2)
    
    testing_tasks = [
        'Test WF1 end-to-end: Send test email → captured in Gmail → ingested to Supabase',
        'Test WF2-5 pipeline: Message classified → context retrieved → draft generated with confidence score',
        'Test WF6 (Gmail only initially): Draft auto-sent via Gmail SMTP',
        'Test WF7: Lead extraction from real message and inserted to Supabase leads table',
        'Test WF8: Dashboard approve/reject/regenerate webhooks trigger correct n8n actions',
        'Document any issues in tech-debt.md for Week 4 Polish phase'
    ]
    
    for task in testing_tasks:
        doc.add_paragraph(task, style='List Bullet')
    
    doc.add_paragraph('✓ Success Criteria: Full pipeline processed 5 real emails → 5 drafts generated → 0 errors', 
                      style='List Bullet')
    doc.add_paragraph()
    
    # ==================== USAMA KHAN TASKS ====================
    add_styled_heading(doc, '🎨 USAMA KHAN - Dashboard UI Enhancements (18 hours)', level=1)
    
    usama_heading = doc.add_paragraph('Priority: Can work independently of infrastructure. Use existing Supabase seed data.')
    usama_heading.runs[0].font.italic = True
    
    doc.add_paragraph()
    
    table4 = doc.add_table(rows=1, cols=4)
    table4.style = 'Light Grid Accent 1'
    add_table_row(table4, ['Feature', 'Description', 'Hours', 'Week'], is_header=True)
    
    usama_tasks = [
        ['Regenerate Modal', 'Modal with feedback textarea + submit button + calls /api/drafts/:id/regenerate', '1.5 hrs', 'Week 2'],
        ['Rejection Modal', 'Modal to capture rejection reason → calls /api/drafts/:id/reject with reason', '1 hr', 'Week 2'],
        ['Version History', 'Expandable section showing draft version history with timestamps + who edited', '2 hrs', 'Week 2'],
        ['Context Sources', 'KB source display - show which KB entries influenced draft generation', '1.5 hrs', 'Week 2-3'],
        ['Bulk Actions', 'Multi-select inbox → bulk approve/reject/reassign from context menu', '3 hrs', 'Week 3'],
        ['Analytics Page', 'New /analytics route: approval rate chart, response time stats, category breakdown', '4 hrs', 'Week 3'],
        ['Settings Page', 'New /settings route: tone preferences, auto-send toggles, team member management', '3 hrs', 'Week 3'],
        ['Polish & Testing', 'CSS refinements, accessibility audit (a11y), cross-browser testing', '1.5 hrs', 'Week 4'],
    ]
    
    for task in usama_tasks:
        add_table_row(table4, task)
    
    doc.add_paragraph('✓ Deliverable: All 8 features merged to main branch + all E2E tests passing', 
                      style='List Bullet')
    doc.add_paragraph('✓ Figma Mockups: Share mockups with Sophie before implementation for visual consistency', 
                      style='List Bullet')
    doc.add_paragraph()
    
    # ==================== SOPHIE LAURENT TASKS ====================
    add_styled_heading(doc, '📝 SOPHIE LAURENT - Brand Voice & Content (30 hours) - CRITICAL PATH', level=1)
    
    critical_note = doc.add_paragraph('⚠️ CRITICAL: This work is blocking all prompt finalization. START IMMEDIATELY. All phases depend on this.')
    critical_note.runs[0].font.bold = True
    critical_note.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph()
    
    table5 = doc.add_table(rows=1, cols=4)
    table5.style = 'Light Grid Accent 1'
    add_table_row(table5, ['Phase', 'Task', 'Hours', 'Week'], is_header=True)
    
    sophie_tasks = [
        ['Phase 1: Brand Voice', '1.1 Email 5-10 brand voice example emails to team', '4 hrs', 'Week 2'],
        ['', '1.2 Finalize P1 brand voice system prompt (incorporates feedback)', '2 hrs', 'Week 2'],
        ['Phase 2: Templates', '2.1 Provide 15+ response templates by category (complaint, inquiry, order, support)', '4 hrs', 'Week 2'],
        ['', '2.2 Create tone guidelines (professional, empathetic, urgent, promotional)', '2 hrs', 'Week 2'],
        ['Phase 3: KB Collection', '3.1 Export 100+ historic emails from Paris Dream brand email account', '8 hrs', 'Week 2-3'],
        ['', '3.2 Create 50+ Q&A database (customer FAQs, common issues)', '6 hrs', 'Week 3'],
        ['Phase 4: KB Refinement', '4.1 Bulk upload KB entries to Supabase + trigger WF3 embeddings', '1 hr', 'Week 3'],
        ['', '4.2 Review 20+ auto-generated drafts + calibrate classification model', '3 hrs', 'Week 3-4'],
        ['Phase 5: Final QA', '5.1 Final content audit + sign-off', '2 hrs', 'Week 4'],
    ]
    
    for task in sophie_tasks:
        add_table_row(table5, task)
    
    doc.add_paragraph('✓ Deliverable: P1 prompt finalized + 150+ KB entries populated + 20+ drafts validated', 
                      style='List Bullet')
    doc.add_paragraph('⚠️ Dependencies:', style='List Paragraph')
    doc.add_paragraph('- Blocks: P2-P6 prompts (classification, extraction, tone validation)', style='List Bullet 2')
    doc.add_paragraph('- Blocks: WF5 draft generation quality', style='List Bullet 2')
    doc.add_paragraph('✓ Action: Email brand voice guide to Abdur by end of Week 2 so prompts lock in while infrastructure builds', 
                      style='List Bullet')
    doc.add_paragraph()
    
    # ==================== OPTIONAL DEVOPS TASKS ====================
    add_styled_heading(doc, '🚀 OPTIONAL: DevOps & Monitoring (10 hours)', level=1)
    
    optional_note = doc.add_paragraph('Optional for Week 3-4 (post-MVP). Recommend for production deployment.')
    optional_note.runs[0].font.italic = True
    
    doc.add_paragraph()
    
    table6 = doc.add_table(rows=1, cols=4)
    table6.style = 'Light Grid Accent 1'
    add_table_row(table6, ['Task', 'Description', 'Hours', 'Week'], is_header=True)
    
    devops_tasks = [
        ['Monitoring Setup', 'Prometheus + Grafana or Datadog dashboard (CPU, memory, API latency)', '3.5 hrs', 'Week 3-4'],
        ['Backup Strategy', 'Automated daily snapshots: n8n data + Supabase backups to S3', '2 hrs', 'Week 3-4'],
        ['CI/CD Pipeline', 'GitHub Actions: auto-deploy on main branch → restart backend + rebuild dashboard', '3 hrs', 'Week 4'],
        ['Runbook', 'Troubleshooting guide for common issues (n8n crashes, API 500s, stuck workflows)', '1.5 hrs', 'Week 4'],
    ]
    
    for task in devops_tasks:
        add_table_row(table6, task)
    
    doc.add_paragraph('✓ Deliverable: Production-ready monitoring + automated backups + 1-click deploy', 
                      style='List Bullet')
    doc.add_paragraph()
    
    # ==================== CRITICAL DEPENDENCIES & BLOCKERS ====================
    add_styled_heading(doc, '🔴 Critical Dependencies & Status Blockers', level=1)
    
    doc.add_heading('Critical Path Sequence:', level=2)
    critical_path = [
        '1. VPS Infrastructure (Abdur, Week 2) → n8n running',
        '2. Brand Voice Content (Sophie, Week 2) → P1 prompt finalized',
        '3. OAuth Credentials (Abdur, Week 2) → Gmail API connected',
        '4. Workflow Implementation (Abdur, Week 3-4) → all 8 WFs active',
        '5. UI Enhancements (Usama, Weeks 2-4) → dashboard ready',
        '6. Integration Testing (Abdur, Week 4) → production verification',
    ]
    for i, step in enumerate(critical_path, 1):
        doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph()
    
    doc.add_heading('Blocked Items:', level=2)
    blocked_items = [
        'WhatsApp/Instagram (WF6 partial) - Blocked by Meta App Review (1-4 weeks approval time)',
        'Production Deployment (Week 5+) - Blocked until all workflows tested + KB finalized',
        'Analytics Page (Usama) - Independent, no blockers',
        'Regenerate Modal (Usama) - Depends on WF5 ready (Abdur Week 3-4)',
    ]
    for item in blocked_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    # ==================== WEEKLY MILESTONES ====================
    add_styled_heading(doc, '📊 Weekly Milestones & Go/No-Go Criteria', level=1)
    
    doc.add_heading('Week 2 Milestones (End of April 14):', level=2)
    week2_go = [
        '✓ VPS running + n8n accessible at http://<IP>:5678',
        '✓ All credential nodes created in n8n',
        '✓ Gmail OAuth flow tested + working',
        '✓ P1 Brand Voice Prompt finalized + sent to Abdur',
        '✓ 50+ KB entries staged for bulk upload',
        '✓ UI mockups approved (Usama → Sophie)',
        'GO SIGNAL: All above ✓ → Proceed to Week 3 workflows',
    ]
    for item in week2_go:
        p = doc.add_paragraph(item)
        if '✓' in item:
            p.style = 'List Bullet'
        else:
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_paragraph()
    
    doc.add_heading('Week 3 Milestones (End of April 21):', level=2)
    week3_go = [
        '✓ WF1-WF5 fully implemented in n8n',
        '✓ Email ingestion pipeline end-to-end tested (5 real emails)',
        '✓ Draft generation + tone scoring validated',
        '✓ 150+ KB entries embedded in Pinecone',
        '✓ Regenerate + Rejection modals merged to main',
        '✓ Version history view deployed',
        'GO SIGNAL: All above ✓ → Proceed to Week 4 finalization',
    ]
    for item in week3_go:
        p = doc.add_paragraph(item)
        if '✓' in item:
            p.style = 'List Bullet'
        else:
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_paragraph()
    
    doc.add_heading('Week 4 Milestones (End of April 25 - LAUNCH):', level=2)
    week4_go = [
        '✓ WF6-WF8 fully implemented + tested',
        '✓ Full 8-workflow pipeline end-to-end verified',
        '✓ Step 14 verification checklist all ✓',
        '✓ All dashboard UI features live + tested',
        '✓ 20+ auto-draft samples reviewed + calibrated',
        '✓ Production runbook + monitoring dashboard live',
        '✓ All 103 hours of pending work completed',
        'GO SIGNAL: LAUNCH DPCA to production 🚀',
    ]
    for item in week4_go:
        p = doc.add_paragraph(item)
        if '✓' in item:
            p.style = 'List Bullet'
        else:
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ==================== HANDOFF INSTRUCTIONS ====================
    add_styled_heading(doc, '📋 Handoff Instructions', level=1)
    
    doc.add_heading('For Abdur Rehman:', level=2)
    doc.add_paragraph('Week 2: Set up VPS + n8n + credentials', style='List Number')
    doc.add_paragraph('Week 3-4: Implement all 8 n8n workflows from JSON stubs', style='List Number')
    doc.add_paragraph('Setup checklist: docs/SETUP.md (VPS deployment guide)', style='List Number')
    doc.add_paragraph('Workflow templates: n8n-workflows/ folder (8 JSON files)', style='List Number')
    
    doc.add_paragraph()
    
    doc.add_heading('For Usama Khan:', level=2)
    doc.add_paragraph('Start immediately: UI mockups from Figma', style='List Number')
    doc.add_paragraph('Week 2: Regenerate modal + rejection modal + version history', style='List Number')
    doc.add_paragraph('Week 3: Bulk actions + analytics + settings pages', style='List Number')
    doc.add_paragraph('Component stubs: dashboard/app/components/ (placeholder files ready)', style='List Number')
    
    doc.add_paragraph()
    
    doc.add_heading('For Sophie Laurent (CRITICAL):', level=2)
    doc.add_paragraph('CRITICAL: Start Week 2, do NOT wait', style='List Number')
    doc.add_paragraph('Email brand voice examples + 15+ templates by April 10', style='List Number')
    doc.add_paragraph('Finalize P1 system prompt by April 12', style='List Number')
    doc.add_paragraph('Export 100+ historic emails + create 50+ Q&A by April 17', style='List Number')
    doc.add_paragraph('Content templates: docs/BRAND_VOICE.md + docs/TEMPLATES.md', style='List Number')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer_para = doc.add_paragraph('──────────────────────────────────────────')
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    footer = doc.add_paragraph('End of Sprint Roadmap Document')
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer.runs[0].font.italic = True
    footer_para2 = doc.add_paragraph('For questions or changes, update this document and redistribute to all team members.')
    footer_para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_para2.runs[0].font.size = Pt(10)
    
    return doc

if __name__ == '__main__':
    doc = create_document()
    output_path = r'd:\workspace\DPCA\Pending_Tasks_Sprint_Roadmap.docx'
    doc.save(output_path)
    print(f'✓ Document created successfully: {output_path}')
    print(f'✓ File size: {len(open(output_path, "rb").read()) / 1024:.1f} KB')
    print('✓ Ready to distribute to team: Abdur, Usama, Sophie')
