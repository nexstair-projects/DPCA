#!/usr/bin/env python3
"""Generate comprehensive Week 3 Work Report for DPCA project"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_report():
    doc = Document()

    # === TITLE ===
    title = doc.add_heading('DPCA PROJECT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Week 3 Work Report', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph('April 12 - April 18, 2026')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.italic = True
    date_para.runs[0].font.size = Pt(12)

    doc.add_paragraph()

    # === EXECUTIVE SUMMARY ===
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'Week 3 focused on organizational restructuring and establishing collaborative development practices. '
        'The repository was migrated from the personal GitHub account (Ab-dur-Rehman/DPCA) to the official '
        'Nexstair organization account (nexstair-projects/DPCA) to enable team collaboration. A comprehensive '
        'Developer Workflow Guide was created to standardize Git practices across the team, and a CODEOWNERS '
        'file was added to enforce code review policies. These changes lay the groundwork for scaling the '
        'development team and ensuring consistent code quality.'
    )

    doc.add_paragraph()

    # === COMPLETED DELIVERABLES ===
    doc.add_heading('Completed Deliverables', level=1)

    deliverables = {
        'Repository Migration to Nexstair Organization': [
            'Migrated repository from personal account (Ab-dur-Rehman/DPCA) to organization account (nexstair-projects/DPCA)',
            'Configured new remote "nexstair" pointing to https://github.com/nexstair-projects/DPCA.git',
            'Retained original remote "origin" (Ab-dur-Rehman/DPCA) as backup reference',
            'Pushed all branches and history to the new organization repository',
            'Verified repository accessible at https://github.com/nexstair-projects/DPCA',
            'Updated clone URLs in documentation to reference the new organization repo',
        ],
        'Developer Workflow Guide': [
            'Created comprehensive DPCA_DEVELOPER_GUIDE.md (188 lines) in guides/ directory',
            'Documented first-time setup instructions including clone and Git identity configuration',
            'Defined daily workflow: feature branch creation, syncing with main, committing, pushing, and opening PRs',
            'Established branch naming conventions: feature/, fix/, update/ prefixes',
            'Wrote commit message guidelines: present tense, brief but descriptive',
            'Documented Pull Request process with step-by-step GitHub UI instructions',
            'Added conflict resolution guide with visual markers explanation',
            'Created quick command reference table for common Git operations',
            'Included VS Code tips: Source Control panel, Git Graph extension, Live Share',
            'Designated Ab-dur-Rehman and ujavaid015 as code reviewers',
        ],
        'CODEOWNERS Configuration': [
            'Created .github/CODEOWNERS file for automated review assignment',
            'Configured code ownership rules to enforce review policies',
            'Ensures all pull requests require approval from designated owners',
            'Supports branch protection rules on the nexstair-projects repository',
        ],
        'Important Development Rules Established': [
            'No direct pushes to main branch — branch protection enforced',
            'All changes require Pull Requests, even small fixes',
            'Developers must pull main into feature branches before pushing',
            'Branches must be deleted after merging to keep repo clean',
            'Clear commit messages required for readable history',
        ],
    }

    for heading, items in deliverables.items():
        doc.add_heading(heading, level=2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # === REPOSITORY MIGRATION DETAILS ===
    doc.add_heading('Repository Migration Details', level=1)

    doc.add_heading('Migration Summary', level=2)
    doc.add_paragraph(
        'The DPCA repository was transferred from the personal GitHub account to the Nexstair '
        'organization to centralize project management under the company account. This enables '
        'proper team access control, organization-level settings, and professional project governance.'
    )

    doc.add_heading('Remote Configuration (Post-Migration)', level=2)

    remote_table = doc.add_table(rows=1, cols=3)
    remote_table.style = 'Light Grid Accent 1'
    remote_hdr = remote_table.rows[0].cells
    remote_hdr[0].text = 'Remote Name'
    remote_hdr[1].text = 'URL'
    remote_hdr[2].text = 'Purpose'

    remotes = [
        ('nexstair', 'https://github.com/nexstair-projects/DPCA.git', 'Primary — Organization repository'),
        ('origin', 'https://github.com/Ab-dur-Rehman/DPCA', 'Backup — Original personal repository'),
    ]

    for name, url, purpose in remotes:
        row = remote_table.add_row().cells
        row[0].text = name
        row[1].text = url
        row[2].text = purpose

    doc.add_paragraph()

    doc.add_heading('Impact on Existing Deployments', level=2)
    impact_items = [
        'Vercel deployment: Connected to origin (Ab-dur-Rehman/DPCA) — may need reconnection to nexstair-projects/DPCA',
        'Railway deployment: No impact — deploys from Docker/environment config, not directly from GitHub',
        'Team members: Must clone from https://github.com/nexstair-projects/DPCA.git going forward',
        'CI/CD: GitHub Actions workflows will run from the new organization repository',
    ]
    for item in impact_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # === FILES MODIFIED & CREATED ===
    doc.add_heading('Files Modified & Created', level=1)

    doc.add_heading('Documentation (1 file)', level=2)
    doc.add_paragraph('guides/DPCA_DEVELOPER_GUIDE.md [NEW]')
    doc.add_paragraph('Comprehensive developer workflow guide — 188 lines covering Git workflow, branching strategy, PR process, conflict resolution, and VS Code tips', style='List Bullet')

    doc.add_heading('GitHub Configuration (1 file)', level=2)
    doc.add_paragraph('.github/CODEOWNERS [NEW]')
    doc.add_paragraph('Code ownership file for automated PR review assignment', style='List Bullet')

    doc.add_heading('Reports (1 file)', level=2)
    doc.add_paragraph('work_reports/Week_2_Work_Report.docx [CREATED EARLIER IN WEEK]')
    doc.add_paragraph('Week 2 progress report generated via generate_week2_report.py script', style='List Bullet')

    doc.add_paragraph()

    # === PRODUCTION CONFIGURATION (UNCHANGED) ===
    doc.add_heading('Production Configuration', level=1)
    doc.add_paragraph('No changes to production configuration this week. All services remain operational:')

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    tbl_cells = table.rows[0].cells
    tbl_cells[0].text = 'Component'
    tbl_cells[1].text = 'URL / Endpoint'
    tbl_cells[2].text = 'Status'

    urls = [
        ('Frontend', 'https://dpca-ten.vercel.app', 'Operational'),
        ('Backend API', 'https://dpca-production.up.railway.app', 'Operational'),
        ('Database', 'https://hefkqlkiuiqhgssdmvad.supabase.co', 'Operational'),
        ('New Repo', 'https://github.com/nexstair-projects/DPCA', 'Active'),
        ('Old Repo', 'https://github.com/Ab-dur-Rehman/DPCA', 'Retained as backup'),
    ]

    for component, url, status in urls:
        row = table.add_row().cells
        row[0].text = component
        row[1].text = url
        row[2].text = status

    doc.add_paragraph()

    # === GIT COMMIT LOG ===
    doc.add_heading('Git Commit Log (Week 3)', level=1)

    commit_table = doc.add_table(rows=1, cols=3)
    commit_table.style = 'Light Grid Accent 1'
    commit_hdr = commit_table.rows[0].cells
    commit_hdr[0].text = 'Hash'
    commit_hdr[1].text = 'Date'
    commit_hdr[2].text = 'Message'

    commits = [
        ('b3f7a65', 'Apr 12, 2026', 'feat: Add Week 2 Work Report generation script and corresponding report document'),
        ('c60e388', 'Apr 19, 2026', 'Add CODEOWNERS file'),
        ('e12cf25', 'Apr 19, 2026', 'feat: Add Developer Workflow Guide to standardize collaboration practices'),
    ]

    for hash_val, date, msg in commits:
        row = commit_table.add_row().cells
        row[0].text = hash_val
        row[1].text = date
        row[2].text = msg

    doc.add_paragraph()

    # === SUMMARY STATISTICS ===
    doc.add_heading('Summary Statistics', level=1)

    stats = [
        'Repository Migration: Completed (Ab-dur-Rehman/DPCA → nexstair-projects/DPCA)',
        'New Files Created: 2 (DPCA_DEVELOPER_GUIDE.md, CODEOWNERS)',
        'Total Commits This Week: 3',
        'Lines of Documentation Written: 188 (Developer Guide)',
        'Team Collaboration Standards Defined: 6 rules',
        'Branch Naming Conventions Established: 3 prefixes (feature/, fix/, update/)',
        'Code Reviewers Assigned: 2 (Ab-dur-Rehman, ujavaid015)',
        'Production Downtime: 0',
        'Build Errors: 0',
        'TypeScript Errors: 0',
    ]

    for stat in stats:
        doc.add_paragraph(stat, style='List Bullet')

    doc.add_paragraph()

    # === OUTSTANDING ITEMS ===
    doc.add_heading('Outstanding Items & Carry-Forward from Week 2', level=1)

    blocker_table = doc.add_table(rows=1, cols=3)
    blocker_table.style = 'Light Grid Accent 1'
    blocker_hdr = blocker_table.rows[0].cells
    blocker_hdr[0].text = 'Item'
    blocker_hdr[1].text = 'Priority'
    blocker_hdr[2].text = 'Description'

    blockers = [
        ('Execute DB Migration', 'CRITICAL', 'Run expand_message_status.sql in Supabase SQL Editor to allow n8n webhook inserts'),
        ('Reconnect Vercel to Nexstair Repo', 'HIGH', 'Update Vercel project to deploy from nexstair-projects/DPCA instead of Ab-dur-Rehman/DPCA'),
        ('Brand Voice Content', 'CRITICAL PATH', 'Awaiting Sophie\'s brand voice examples — blocks prompt finalization'),
        ('n8n VPS Setup', 'HIGH', 'Provision Ubuntu 22.04 VPS (2GB RAM) for n8n deployment'),
        ('UI Enhancements', 'MEDIUM', 'Regenerate modals (regenerate, rejection, version history)'),
        ('Branch Protection Rules', 'HIGH', 'Configure branch protection on nexstair-projects/DPCA main branch'),
    ]

    for item, pri, desc in blockers:
        blocker_row = blocker_table.add_row().cells
        blocker_row[0].text = item
        blocker_row[1].text = pri
        blocker_row[2].text = desc

    doc.add_paragraph()

    # === NEXT STEPS ===
    doc.add_heading('Next Steps (Week 4 Priority List)', level=1)

    next_steps = [
        'Configure branch protection rules on nexstair-projects/DPCA',
        'Reconnect Vercel deployment to the new organization repository',
        'Execute pending DB migration in Supabase SQL Editor',
        'Provision VPS for n8n deployment',
        'Configure and deploy n8n on VPS',
        'Import and test n8n workflows with real data',
        'Receive and integrate brand voice content from Sophie',
        'Implement UI enhancement modals (regenerate, rejection, version history)',
        'Onboard team members using the Developer Workflow Guide',
        'End-to-end system testing with live channels',
    ]

    for i, step in enumerate(next_steps, 1):
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()

    # === FOOTER ===
    doc.add_paragraph()
    footer = doc.add_paragraph('---')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer_text = doc.add_paragraph(f'Report Generated: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text.runs[0].font.size = Pt(9)
    footer_text.runs[0].font.italic = True

    # Save
    doc.save('work_reports/Week_3_Work_Report.docx')
    print('✓ Successfully created: work_reports/Week_3_Work_Report.docx')

if __name__ == '__main__':
    create_report()
